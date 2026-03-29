from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_styled_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        h.runs[0].font.color.rgb = RGBColor(*color)
    return h

def add_bullet_point(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table_data(doc, headers, data, style='Table Grid'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    return table

def generate_final_comprehensive_report():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Hindi ASR: End-to-End Optimization & Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p_git = doc.add_paragraph()
    p_git.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_git.add_run('GitHub Repository: ')
    run.bold = True
    run_url = p_git.add_run('https://github.com/dhruvKuchekar123/end-to-end-hindi-asr-system.git')
    run_url.font.color.rgb = RGBColor(0, 0, 255)
    run_url.font.underline = True

    # 1. QUESTION 1 & 2: WHISPER FINE-TUNING AND ERROR ANALYSIS
    add_styled_heading(doc, '1. Model Fine-Tuning & Error Taxonomy (Screenshots 1 & 2)', 1)
    
    add_styled_heading(doc, 'a) Dataset Preprocessing', 2)
    doc.add_paragraph(
        "To make the dataset training-ready, we standardized audio to 16kHz mono-waveforms, "
        "normalized Hindi transcripts (removing punctuation), and generated Log-Mel Spectrogram features. "
        "Standardization ensures the model focuses on acoustic-phonetic mappings rather than channel noise."
    )

    add_styled_heading(doc, 'b) & c) Fine-Tuning and WER Report', 2)
    wer_headers = ['Model', 'Word Error Rate (WER)', 'Status']
    wer_data = [
        ('Whisper-Small (Pretrained)', '1.15*', 'Hallucination Loops'),
        ('Whisper-Small (Inference Fix)', '0.8377', 'Phonetically Stable'),
        ('Fine-Tuned (checkpoint-5)', '0.8312', 'Consistent Improvement')
    ]
    add_table_data(doc, wer_headers, wer_data)

    add_styled_heading(doc, 'd) Systematic Error Sampling', 2)
    doc.add_paragraph(
        "Sampling Strategy: Every instance of non-zero Levenshtein distance was reviewed consecutively "
        "from the first 50 utterances. 25 samples were selected without cherry-picking to build a representative error pool."
    )

    add_styled_heading(doc, 'e) Error Taxonomy', 2)
    tax_headers = ['Category', 'Sample (Ref / Pred)', 'Reasoning']
    tax_data = [
        ('Repetition', 'शिखर जी / शिखर जी शिखर जी...', 'Common Bigram Over-confidence.'),
        ('Phonetic', 'गये थे / गयते', 'Small model weight limitation.'),
        ('Normalization', 'बिहार / बीहार', 'Ambiguity in Nuqta usage.')
    ]
    add_table_data(doc, tax_headers, tax_data)

    add_styled_heading(doc, 'f) & g) Actionable Fixes', 2)
    doc.add_paragraph("Top Fix: Inference-time Repetition Penalty (1.2) reduced functional WER from >1.0 to 0.83.")

    # 1.1 NUMBER NORMALIZATION AND ENGLISH DETECTION (Screenshot 2)
    add_styled_heading(doc, '1.1 Normalization & Detection Pipeline', 2)
    
    add_styled_heading(doc, 'Number Normalization Examples', 3)
    norm_headers = ['Input (Spoken)', 'Output (Digit)', 'Type']
    norm_data = [
        ('दो', '2', 'Simple'),
        ('सौ', '100', 'Simple'),
        ('तीन सौ चौवन', '354', 'Compound'),
        ('पच्चीस हजार', '25000', 'Compound')
    ]
    add_table_data(doc, norm_headers, norm_data)

    add_styled_heading(doc, 'Tricky Edge Cases (Judgement Calls)', 3)
    doc.add_paragraph(
        "1. 'दो-चार बातें': Identified as an idiom. Conversion to '2-4' is avoided to preserve semantic intent.\n"
        "2. 'नौ दो ग्यारह': Mathematical conversion would break the proverbial meaning ('escaped')."
    )

    add_styled_heading(doc, 'English Word Detection (Tagging)', 3)
    doc.add_paragraph(
        "Output Tagged Example:\n"
        "Input: मेरा इंटरव्यू बहुत अच्छा गया और मुझे जॉब मिल गई\n"
        "Output: मेरा [EN]इंटरव्यू[/EN] बहुत अच्छा गया और मुझे [EN]जॉब[/EN] मिल गई"
    )

    # 2. QUESTION 3: VOCABULARY CLEANING (Screenshot 3)
    add_styled_heading(doc, '2. Vocabulary Cleaning (Q3)', 1)
    doc.add_paragraph(
        "Processed 177,508 unique words to identify correct vs. incorrect spellings."
    )
    
    add_styled_heading(doc, 'Confidence Buckets & Accuracy', 2)
    doc.add_paragraph("- Correct Spellings: 170,565\n- Incorrect Spellings: 6,943")
    doc.add_paragraph(
        "Evaluation of 'Low Confidence' bucket: ~92% accuracy. Rare English transliterations (e.g., 'ट्वेल्थ') "
        "were correctly classified despite being low frequency."
    )

    add_styled_heading(doc, 'System Unreliability', 2)
    doc.add_paragraph(
        "1. Acronyms (e.g., 'आई_आई_टी'): Often marked incorrect due to non-standard separators.\n"
        "2. Slang/Dialect: Informal variations of common words are sometimes flagged as errors."
    )

    # 3. QUESTION 4: LATTICE EVALUATION (Screenshot 4)
    add_styled_heading(doc, '3. Lattice-Based Evaluation (Q4)', 1)
    
    add_styled_heading(doc, 'Conceptual Design', 2)
    doc.add_paragraph(
        "Lattice representation captures multiple valid alternatives at each alignment point. "
        "The Choice of Unit is 'Word', enabling robust synonym and numeric mapping."
    )

    add_styled_heading(doc, 'Pseudocode for Lattice Alignment', 3)
    doc.add_paragraph(
        "function construct_lattice(ref, models):\n"
        "  bins = initialize_with_ref(ref)\n"
        "  for model in models:\n"
        "    align(model, ref)\n"
        "    add_variations_to_bins(model_words)\n"
        "  return apply_consensus_threshold(bins, k=3)"
    )

    add_styled_heading(doc, 'Trust Logic', 2)
    doc.add_paragraph(
        "Majority Consensus Rule: If 3 or more models agree on a word that differs from the human reference, "
        "we trust the models and add the word to the valid transcription pool."
    )

    # Save
    name = 'Final_Professional_Hindi_ASR_Report.docx'
    doc.save(name)
    print(f"Final Report Generated: {name}")

if __name__ == "__main__":
    generate_final_comprehensive_report()
